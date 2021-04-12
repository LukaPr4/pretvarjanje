from tkinter import *
from tkinter import ttk
from tkinter import filedialog
from os import path
from docx import Document
from bs4 import BeautifulSoup
import mammoth
import csv
import html2text
import codecs
import pandas as pd
import io

class Izgled(object): #uporabniški vmesnik
    def __init__(self, master):
        self.doc = Document()
        self.extention = StringVar()
        self.extention2 = StringVar()
        self.filename = StringVar()
        
        self.master = master
        self.levaStran()
        self.desnaStran()
        self.pretvori()
        
    def levaStran(self): #gumb, kobinacijsko oknce in napis na levi strani programa
        self.label1 = Label(self.master, text = "IZ")
        self.label1.grid(column = 0, row = 0)

        self.button = ttk.Button(self.master, text = "Poiščite datoteko", command = self.snapis)
        self.button.grid(column = 0, row = 2)    

        self.izbira1 = ttk.Combobox(self.master, width = 16, values = [".txt", ".docx",".csv", ".html"], state= 'readonly', textvariable = self.extention)
        self.izbira1.grid(column = 0, row = 1)

    def desnaStran(self): #napis in kobinacijsko oknce na desni strani programa
        self.label2 = Label(self.master, text = "V")
        self.label2.grid(column = 3, row = 0)

        self.izbira2 = ttk.Combobox(self.master, width = 16, values = [".txt", ".docx",".csv", ".html"], state= 'readonly', textvariable = self.extention2)
        self.izbira2.grid(column = 3, row = 1)

    def snapis(self):
        self.snapis = Label(self.master, text = " ") #napis kjer se pojavi ime izbrane datoteke
        self.filename = ""
        self.filename = filedialog.askopenfilename(initialdir = "/", title = ("Izberite", self.extention.get() , "Datoteko"), filetype = (("Datoteka", self.extention.get()), ("All Files", "*.*")))#odpre oknce kjer uporabnik izbere datoteko
        self.filename2 = path.basename(self.filename) #ime izbrane datoteke
        self.snapis.configure(text = self.filename2)
        self.snapis.grid(column = 0, row =4) #postavi napis na določeno mesto v oknu
        
    def pretvori(self):
        self.convert = Button(self.master, text= "pretvori", command = self.pretvornik) #gumb z vgrajeno funkcijo za pretvarjanje
        self.convert.grid(column = 2, row = 1) #postavljen na določeno mesto v oknu


    def pretvornik(self): #poskrbi da ima vsaka kombinacija vnosov v kobinacijsko oknce rezultat
        if self.extention.get() == ".txt": #možnost vnosa prvega oknca
            if self.extention2.get() == ".docx": #možnost vnosa drugega oknca
                Funkcije.fromTxtToDocx(self)
            elif self.extention2.get() == ".csv":
                Funkcije.fromTxtToCsv(self)
            elif self.extention2.get() == ".html":
                Funkcije.fromTxtToHtml(self)

        elif self.extention.get() == ".docx":
            if self.extention2.get() == ".txt":
                Funkcije.fromDocxToTxt(self)
            elif self.extention2.get() == ".csv":
                Funkcije.fromDocxToCsv(self)
            elif self.extention2.get() == ".html":
                Funkcije.fromDocxToHtml(self)

        elif self.extention.get() == ".csv":
            if self.extention2.get() == ".txt":
                Funkcije.fromCsvToTxt(self)
            elif self.extention2.get() == ".docx":
                Funkcije.fromCsvToDocx(self)
            elif self.extention2.get() == ".html":
                Funkcije.fromCsvToHtml(self)

        else:
            if self.extention2.get() == ".txt":
                Funkcije.fromHtmlToTxt(self)
            elif self.extention2.get() == ".csv":
                Funkcije.fromHtmlToCsv(self)
            elif self.extention2.get() == ".docx":
                Funkcije.fromHtmlToDocx(self)
        
    


class Funkcije(object):
    def fromTxtToDocx(self): #prevede .txt datoteko v .docx datoteko
        self.podatki = ""
        with codecs.open(self.filename, "r", encoding="utf8") as f: #odpre datoteko s pomočjo knjižnice codecs, ki omogoča UTF-8
            self.podatki = f.read()
            self.doc.add_paragraph(self.podatki) #napiše vse podatke v .docx datoteko
        self.doc.save(str(path.splitext(self.filename2)[0] + ".docx")) #shrani z enakim imenom datoteke kot ga je imela prejšnja, vendar z drugo končnico

    def fromDocxToTxt(self): #prevede .docx datoteko v .txt
        self.doc2 = Document(self.filename) #nov dokument
        with codecs.open(str(path.splitext(self.filename2)[0] + ".txt"), "w", encoding = 'utf-8') as f: #odpre .txt s kodiranje v UTF-8
            for p in self.doc2.paragraphs: #zapiše vsak odstavek posebej v .txt
                f.write(p.text)

    def fromTxtToCsv(self): #prevede .txt v .csv
        self.datoteka = pd.read_csv(self.filename) #prebere .txt
        self.datoteka.to_csv(str(path.splitext(self.filename2)[0] + ".csv"), index = None) #shrani v .csv

    def fromTxtToHtml(self):#.txt v .html
        self.vsebina = codecs.open(self.filename, 'r', encoding="utf8" ) #shrani vsebnio .txt in poskrbi da napačne črke mora zamenjati
        with codecs.open(str(path.splitext(self.filename2)[0] + ".html"), "w", encoding="utf8") as f: #nova .html 
            for p in self.vsebina.readlines(): #zapiše celotno besedilo po vrstah
                f.write("<pre>" + p + "</pre> <br>\n") #poskrbi za formatiran zapis z novimi vrsticami nakoncu vsake vrste

    def fromDocxToCsv(self): #.docx v .csv
        self.doc2 = Document(self.filename) #nov dokument

        self.podatki = ""
        for p in self.doc2.paragraphs: #shrani podatke iz celotnega dokumenta v spremenljivko
            self.podatki += p.text

        with codecs.open(str(path.splitext(self.filename2)[0] + ".csv"), "w", encoding="utf8") as f: #odpre .csv
            for i in self.podatki: #prepiše vse podatke po vrsticah iz spremenljivke
                f.writelines(i)

    def fromDocxToHtml(self): #.docx to .html
        self.f = open(self.filename, 'rb') #odpre .docx v rb, kar pomeni v branje binarnem načinu
        self.b = open(str(path.splitext(self.filename2)[0] + ".html"), 'wb') #odpre .html v wb, kar pomeni pisanje in branje v binarnem načinu
        self.document = mammoth.convert_to_html(self.f) #prevede .docx s pomočjo knjižnice mammoth
        self.b.write(self.document.value.encode('utf8')) #zapiše besedilo v .html
        self.f.close() #zapre datoteki
        self.b.close()

    def fromCsvToTxt(self): #.csv v .txt
        with codecs.open(str(path.splitext(self.filename2)[0] + ".txt"), "w", encoding="utf8") as f: #odpre .txt v UTF-8
            with codecs.open(self.filename2, "r", encoding="utf8") as f1: #odpre .csv
                [f.write(" ". join(r)+"\n") for r in csv.reader(f1, delimiter = ',')]#vsako del tabele ločen z vejicami iz .csv prepiše v .txt
            f.close()

    def fromCsvToDocx(self): #.csv v .docx
        with codecs.open(self.filename, 'r', encoding="utf8") as f: #odpre .csv
            self.csvv = csv.reader(f, delimiter = ',')
            for i in self.csvv: #prebere celotno besedilo
                for k in i: #zapiše besedilo
                    self.doc.add_paragraph(k)

        self.doc.save(str(path.splitext(self.filename2)[0] + ".docx")) #shrani .docx
        
    def fromHtmlToTxt(self): #.html v .txt
        self.f = codecs.open(self.filename, "r", encoding="utf8") #odpre .html
        self.tekst = self.f.read()
        self.besedilo = BeautifulSoup(self.tekst, features='lxml') #shrani besedilo v podatkovni tip BeautifulSoup v načinu lxml, ki omogoča enostavno upravljanje z XML in HTML
        with codecs.open(str(path.splitext(self.filename2)[0] + ".txt"), "w", encoding="utf8") as fi: #odpre .txt
            fi.write(self.besedilo.get_text('\n')) #zapiše besedilo

    def fromHtmlToDocx(self): #.html v .docx
        self.f = codecs.open(self.filename, "r", encoding="utf8",) #odpre .html
        self.tekst = self.f.read()
        self.besedilo = BeautifulSoup(self.tekst, features='lxml') #shrani besedilo v podatkovni tip BeautifulSoup v načinu lxml
        self.doc.add_paragraph(self.besedilo.get_text('\n')) #doda odstavke besedila v dokument
        self.doc.save(str(path.splitext(self.filename2)[0] + ".docx")) #shrani novo .docx datoteko


    def fromCsvToHtml(self): #csv v .html
        self.f = pd.read_csv(self.filename) #prebere .csv
        self.f.to_html(str(path.splitext(self.filename2)[0] + ".html")) #shrani v .html

    def fromHtmlToCsv(self): #.html v .csv
        self.tabela = []
        self.glava = [] 
        self.besedilo = codecs.open(self.filename, 'r', encoding="utf8") #odpre .html
        self.soup = BeautifulSoup(self.besedilo) #shrani besedilo v podatkovni tip BeautifulSoup
        self.header = self.soup.find_all("table")[0].find("tr") #poišče prvo značko table v kateri je značka tr
        
        for i in self.header: #poskusi dobiti besedilo iz značk table in tr
            try:
                self.glava.append(i.get_text()) #če so jih shrani v seznam
            except:
                continue
        self.html1 = self.soup.find_all("table")[0].find_all("tr")[1:] #poišče vse naslednje značke table
        for k in self.html1:
            self.html2 = []
            for h in k: #prepiše celotno tabelo v seznam
                try:
                    self.html2.append(h.get_text())
                except:
                    continue
            self.tabela.append(self.html2) 

        self.df = pd.DataFrame(data = self.tabela, columns = self.glava) #shrani tabelo kot podatkovni tip DataFrame

        self.df.to_csv(str(path.splitext(self.filename2)[0] + ".csv")) #iz DataFrama pretvori in shrani v .csv
            
        
root = Tk() #ustvari okno uporabniškega vmesnika
root.geometry("400x340+800+300") #velikost uporabniškega vmesnika
root.title("Prevajanje datotek") #ime celotnega okna
izgled = Izgled(root)
root.mainloop()
